from flask import Blueprint, request, jsonify, send_from_directory
import os, io, json
from datetime import datetime
from PIL import Image, ImageOps
import numpy as np
import cv2
import pytesseract
from docx import Document
from docx.shared import Inches

restore_bp = Blueprint('restore_bp', __name__)

def deskew_image_cv(img):
    coords = cv2.findNonZero(cv2.bitwise_not(img))
    if coords is None:
        return img
    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle
    (h, w) = img.shape[:2]
    center = (w // 2, h // 2)
    M = cv2.getRotationMatrix2D(center, angle, 1.0)
    rotated = cv2.warpAffine(img, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
    return rotated

def restore_pipeline(image_path):
    img = cv2.imread(image_path)
    if img is None:
        raise RuntimeError("failed to read image: " + image_path)

    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
    cl = clahe.apply(gray)

    den = cv2.bilateralFilter(cl, d=9, sigmaColor=75, sigmaSpace=75)

    th = cv2.adaptiveThreshold(den, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                               cv2.THRESH_BINARY, 15, 8)

    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1,1))
    opened = cv2.morphologyEx(th, cv2.MORPH_OPEN, kernel, iterations=1)

    desk = deskew_image_cv(opened)

    restored_rgb = cv2.cvtColor(desk, cv2.COLOR_GRAY2BGR)

    return restored_rgb

def save_restored_and_docx(input_path, restored_img, ocr_text=None, out_dir="data/output/restored", ts=None):
    os.makedirs(out_dir, exist_ok=True)
    if ts is None:
        ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    base = os.path.splitext(os.path.basename(input_path))[0]
    restored_name = f"{base}_restored_{ts}.png"
    restored_path = os.path.join(out_dir, restored_name)

    cv2.imwrite(restored_path, restored_img)

    doc = Document()
    doc.add_heading(f"Restoration Result - {base}", level=1)
    doc.add_paragraph(f"Source: {input_path}")
    doc.add_paragraph(f"Generated: {ts}")
    
    try:
        doc.add_picture(restored_path, width=Inches(6))
    except Exception:
        doc.add_paragraph("[Image embed failed: " + restored_name + "]")

    doc.add_heading("OCR Text (raw)", level=2)
    if ocr_text:
        doc.add_paragraph(ocr_text)
    else:
        doc.add_paragraph("[No OCR text available]")

    docx_name = f"{base}_restored_{ts}.docx"
    docx_path = os.path.join(out_dir, docx_name)
    doc.save(docx_path)

    web_restored = os.path.join("data/output/restored", restored_name)
    web_docx = os.path.join("data/output/restored", docx_name)
    return web_restored, web_docx


@restore_bp.route('/run', methods=['POST'])
def run_restore():
    data = request.get_json() or {}
    image_paths = data.get("images")
    do_ocr = bool(data.get("ocr", True))

    if not image_paths:
        return jsonify({"error": "missing images"}), 400
    if isinstance(image_paths, str):
        image_paths = [image_paths]

    results = []
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

    for img_path in image_paths:
        full = os.path.join(os.getcwd(), img_path)
        try:
            restored = restore_pipeline(full)

            ocr_text = None
            if do_ocr:
                pil = Image.fromarray(cv2.cvtColor(restored, cv2.COLOR_BGR2RGB))
                ocr_text = pytesseract.image_to_string(pil, lang="eng")

            web_img, web_docx = save_restored_and_docx(img_path, restored, ocr_text, ts=ts)

            results.append({
                "image": img_path,
                "restored_image": web_img,
                "docx": web_docx,
                "ocr_text": ocr_text or ""
            })
        except Exception as e:
            results.append({
                "image": img_path,
                "error": str(e)
            })

    return jsonify({"results": results})